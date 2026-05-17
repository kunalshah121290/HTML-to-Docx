from flask import Flask, request, send_file, render_template
from bs4 import BeautifulSoup
from docx import Document
import os
import zipfile
import uuid

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def html_to_docx(html_content, output_file):
    doc = Document()
    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.find_all():
        if element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'li':
            doc.add_paragraph(element.text, style='List Bullet')

    doc.save(output_file)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    try:
        files = request.files.getlist('files')

        if not files:
            return "No files uploaded"

        unique_id = str(uuid.uuid4())
        zip_path = os.path.join(OUTPUT_FOLDER, f"{unique_id}.zip")

        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file in files:
                if file.filename.endswith(".html") or file.filename.endswith(".htm"):
                    
                    html_content = file.read().decode('utf-8')
                    
                    output_filename = file.filename.rsplit('.', 1)[0] + ".docx"
                    output_path = os.path.join(OUTPUT_FOLDER, output_filename)

                    html_to_docx(html_content, output_path)
                    zipf.write(output_path, output_filename)

        return send_file(zip_path, as_attachment=True)

    except Exception as e:
        return f"Error: {str(e)}"

import os
if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 10000)))
